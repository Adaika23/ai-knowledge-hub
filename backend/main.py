from datetime import datetime
import os
import math
import fitz
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from fastapi import UploadFile, File

from fastapi.responses import StreamingResponse

from fastapi import FastAPI, HTTPException, Depends
from fastapi.security import OAuth2PasswordBearer
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI
from passlib.context import CryptContext
from jose import jwt
from sqlalchemy.orm import Session

from database import engine, get_db
from models import (
    Base,
    Note,
    User,
    Chat,
    Document,
)
from schemas import (
    NoteCreate,
    NoteResponse,
    UserRegister,
    UserLogin,
    AIRequest,
    ChatResponse
)


# ================================
# Load Environment Variables
# ================================
load_dotenv()


# ================================
# Create Database Tables
# ================================
Base.metadata.create_all(bind=engine)


# ================================
# Security Configuration
# ================================
SECRET_KEY = os.getenv("SECRET_KEY")

ALGORITHM = os.getenv("ALGORITHM")

ACCESS_TOKEN_EXPIRE_MINUTES = int(
    os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", 30)
)

pwd_context = CryptContext(
    schemes=["bcrypt"],
    deprecated="auto"
)

# ================================
# OAuth2 Token Setup
# ================================
oauth2_scheme = OAuth2PasswordBearer(
    tokenUrl="login"
)

# ================================
# OpenAI Client Setup
# ================================
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ================================
# Generate OpenAI Embedding
# ================================
def generate_embedding(text: str):

    try:

        # Remove invalid Unicode surrogate characters
        text = "".join(
            char
            for char in text
            if not 0xD800 <= ord(char) <= 0xDFFF
        )

        # Convert to safe UTF-8
        text = (
            text
            .encode("utf-8", errors="ignore")
            .decode("utf-8", errors="ignore")
        )

        print("TEXT TYPE:", type(text))
        print("TEXT LENGTH:", len(text))
        print("FIRST 100 CHARACTERS:", repr(text[:100]))

        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )

        # Return embedding vector
        return response.data[0].embedding

    except Exception as e:

        print("Embedding error:", repr(e))

        # Print first invalid character if one exists
        for index, char in enumerate(text):
            code = ord(char)

            if 0xD800 <= code <= 0xDFFF:
                print(
                    f"Invalid surrogate at position {index}: "
                    f"{repr(char)} ({hex(code)})"
                )
                break

        return None
# ================================
# Calculate Cosine Similarity
# ================================
def cosine_similarity(vector_a, vector_b):

    # If either vector is missing, return 0
    if not vector_a or not vector_b:
        return 0

    # Dot product of both vectors
    dot_product = sum(a * b for a, b in zip(vector_a, vector_b))

    # Length of first vector
    magnitude_a = math.sqrt(sum(a * a for a in vector_a))

    # Length of second vector
    magnitude_b = math.sqrt(sum(b * b for b in vector_b))

    # Avoid division by zero
    if magnitude_a == 0 or magnitude_b == 0:
        return 0

    # Cosine similarity score
    return dot_product / (magnitude_a * magnitude_b)

# ================================
# FastAPI App Setup
# ================================
app = FastAPI()


# ================================
# CORS Setup
# ================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Development only
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ================================
# Get Current Logged-in User
# ================================
def get_current_user(
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):

    try:
        # Decode JWT token
        payload = jwt.decode(
            token,
            SECRET_KEY,
            algorithms=[ALGORITHM]
        )

        # Get username from token
        username = payload.get("sub")

        if username is None:
            raise HTTPException(
                status_code=401,
                detail="Invalid authentication token"
            )

    except Exception:
        raise HTTPException(
            status_code=401,
            detail="Invalid authentication token"
        )

    # Find user in database
    db_user = db.query(User).filter(
        User.username == username
    ).first()

    if db_user is None:
        raise HTTPException(
            status_code=401,
            detail="User not found"
        )

    return db_user

# ================================
# Root Route
# ================================
@app.get("/")
def home():
    return {"message": "Backend is running"}


# ================================
# Create Note Route - Protected
# ================================
@app.post("/notes", response_model=NoteResponse)
def create_note(
    note: NoteCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # =====================================
    # Combine Title + Content
    # =====================================
    # This improves embedding quality
    # because OpenAI understands more context.
    full_text = f"{note.title} {note.content}"

    # =====================================
    # Generate OpenAI Embedding Vector
    # =====================================
    embedding = generate_embedding(full_text)

    # =====================================
    # Create Note Object
    # =====================================
    # Save note + embedding + owner ID
    new_note = Note(
        title=note.title,
        content=note.content,
        embedding=embedding,
        user_id=current_user.id
    )

    # =====================================
    # Save Note to PostgreSQL
    # =====================================
    db.add(new_note)
    db.commit()
    db.refresh(new_note)

    # Return saved note
    return new_note


# ================================
# Get All Notes Route - Protected
# ================================
@app.get("/notes", response_model=list[NoteResponse])
def get_notes(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # Get only notes that belong to the logged-in user
    notes = db.query(Note).filter(
        Note.user_id == current_user.id
    ).order_by(Note.id.desc()).all()

    return notes

# ================================
# Delete Note Route - Protected
# ================================
@app.delete("/notes/{note_id}")
def delete_note(
    note_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # Find note that belongs to current user
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    # Note not found
    if note is None:
        raise HTTPException(
            status_code=404,
            detail="Note not found"
        )

    # Delete note
    db.delete(note)

    # Save changes
    db.commit()

    return {
        "message": "Note deleted successfully",
        "deleted_note_id": note_id
    }

# ================================
# Update Note Route - Protected
# ================================
@app.put("/notes/{note_id}", response_model=NoteResponse)
def update_note(
    note_id: int,
    updated_note: NoteCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # Find note that belongs to logged-in user
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    # If note does not exist
    if note is None:
        raise HTTPException(
            status_code=404,
            detail="Note not found"
        )

    # Update note fields
    note.title = updated_note.title
    note.content = updated_note.content

    # Save changes
    db.commit()

    # Refresh updated data
    db.refresh(note)

    return note

# ================================
# Search Notes Route - Protected
# ================================
@app.get("/search", response_model=list[NoteResponse])
def search_notes(
    query: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # Clean search query
    clean_query = query.lower().strip()

    # Search only current user's notes
    results = db.query(Note).filter(
        Note.user_id == current_user.id,
        (
            Note.title.ilike(f"%{clean_query}%") |
            Note.content.ilike(f"%{clean_query}%")
        )
    ).order_by(Note.id.desc()).all()

    return results

# ================================
# Semantic Search Route
# ================================
@app.get("/semantic-search")
def semantic_search(
    query: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    # =====================================
    # Generate embedding for user query
    # =====================================
    query_embedding = generate_embedding(query)

    # If embedding fails
    if not query_embedding:
        raise HTTPException(
            status_code=500,
            detail="Failed to generate query embedding"
        )

    # =====================================
    # Get current user's notes
    # =====================================
    notes = db.query(Note).filter(
        Note.user_id == current_user.id
    ).all()

    # =====================================
    # Calculate similarity scores
    # =====================================
    scored_notes = []

    for note in notes:

        similarity = cosine_similarity(
            query_embedding,
            note.embedding
        )

        scored_notes.append({
            "id": note.id,
            "title": note.title,
            "content": note.content,
            "similarity": similarity
        })

    # =====================================
    # Sort by highest similarity
    # =====================================
    scored_notes.sort(
        key=lambda x: x["similarity"],
        reverse=True
    )

    # Return top 5 most similar notes
    return {
         "results": scored_notes[:5]
    }

# ================================
# 💬 Get Chat History - Protected
# ================================
@app.get("/chat-history", response_model=list[ChatResponse])
def get_chat_history(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chats = db.query(Chat)\
        .filter(Chat.user_id == current_user.id)\
        .order_by(Chat.created_at.asc())\
        .all()

    return chats
# ask - ai
@app.post("/ask-ai")
def ask_ai(
    request: AIRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        today = datetime.now().strftime("%B %d, %Y")

        # Save user question
        user_chat = Chat(
            user_id=current_user.id,
            sender="user",
            message=request.question
        )
        db.add(user_chat)
        db.commit()
        db.refresh(user_chat)

        # Generate question embedding
        question_embedding = generate_embedding(request.question)

        if question_embedding is None:
            msg = "Could not generate embedding."

            ai_chat = Chat(
                user_id=current_user.id,
                sender="assistant",
                message=msg
            )
            db.add(ai_chat)
            db.commit()

            return {
                "answer": msg,
                "sources": []
            }

        # Get user notes and documents
        user_notes = db.query(Note).filter(
            Note.user_id == current_user.id
        ).all()

        user_documents = db.query(Document).filter(
            Document.user_id == current_user.id
        ).all()

        if not user_notes and not user_documents:
            msg = "You do not have any saved notes or uploaded documents yet."

            ai_chat = Chat(
                user_id=current_user.id,
                sender="assistant",
                message=msg
            )
            db.add(ai_chat)
            db.commit()

            return {
                "answer": msg,
                "sources": []
            }

        # Score notes and documents
        scored_items = []

        for note in user_notes:
            if note.embedding:
                scored_items.append({
                    "type": "note",
                    "title": note.title,
                    "content": note.content,
                    "similarity": cosine_similarity(
                        question_embedding,
                        note.embedding
                    )
                })

        for document in user_documents:
            if document.embedding:
                scored_items.append({
                    "type": "document",
                    "title": document.filename,
                    "content": document.content,
                    "similarity": cosine_similarity(
                        question_embedding,
                        document.embedding
                    )
                })

        if not scored_items:
            msg = "No embeddings found in notes or documents."

            ai_chat = Chat(
                user_id=current_user.id,
                sender="assistant",
                message=msg
            )
            db.add(ai_chat)
            db.commit()

            return {
                "answer": msg,
                "sources": []
            }

        # ================================
        # Filter Low Similarity Results
        # ================================
        MIN_SIMILARITY = 0.30

        scored_items.sort(
            key=lambda x: x["similarity"],
            reverse=True
        )

        top_items = [
            item
            for item in scored_items
            if item["similarity"] >= MIN_SIMILARITY
        ][:5]

        # ================================
        # No Relevant Sources Found
        # ================================
        if not top_items:

            answer = "I could not find this in the knowledge base."

            ai_chat = Chat(
                user_id=current_user.id,
                sender="ai",
                message=answer
            )

            db.add(ai_chat)
            db.commit()

            return {
                "answer": answer,
                "sources": []
            }
        # Build knowledge context
        knowledge_context = "\n\n".join([
            f"Source Type: {item['type']}\n"
            f"Title: {item['title']}\n"
            f"Content: {item['content'][:2000]}"
            for item in top_items
        ])

        # Ask OpenAI normally, not streaming
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        f"You are an AI assistant for a personal knowledge hub. "
                        f"Today's date is {today}. "
                        "Answer only using the provided notes and documents. "
                        "Keep answers concise and well organized. "
                        "Use bullet points when helpful. "
                        "Do not copy large sections of documents. "
                        "If the answer is not found, clearly say so."
                    )
                },
                {
                    "role": "user",
                    "content": (
                        f"User question:\n{request.question}\n\n"
                        f"Relevant knowledge base sources:\n{knowledge_context}\n\n"
                        "Answer using clean professional markdown formatting.\n\n"
                        "Rules:\n"
                        "- Use headings when helpful\n"
                        "- Use bullet points for lists\n"
                        "- Use numbered steps for instructions\n"
                        "- Use **bold text** for important points\n"
                        "- Use code blocks for code examples\n"
                        "- Answer only from the knowledge base above\n"
                        "- If the answer is not found, say: I could not find this in the knowledge base."
                    ),
                },
            ]
        )

        answer = response.choices[0].message.content

        # Save AI response
        ai_chat = Chat(
            user_id=current_user.id,
            sender="assistant",
            message=answer
        )
        db.add(ai_chat)
        db.commit()
        db.refresh(ai_chat)

        # Build source cards
        sources = [
            {
                "type": item["type"],
                "title": item["title"],
                "similarity": round(item["similarity"], 4),
                "preview": item["content"][:300]
            }
            for item in top_items
        ]

        return {
            "answer": answer,
            "sources": sources
        }

    except Exception as e:
        return {
            "answer": f"AI error: {str(e)}",
            "sources": []
        }
    
# ================================
# 🗑️ Clear Chat History
# ================================
@app.delete("/chat-history")
def clear_chat_history(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    db.query(Chat).filter(
        Chat.user_id == current_user.id
    ).delete()

    db.commit()

    return {
        "message": "Chat history cleared successfully."
    }  

# ================================
# 📄 Upload Document
# ================================
@app.post("/upload-document")
def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    try:
        filename = file.filename.lower()

        # Read uploaded file bytes
        file_bytes = file.file.read()

        extracted_text = ""

        # ================================
        # Extract Text from TXT
        # ================================
        if filename.endswith(".txt"):
            extracted_text = file_bytes.decode("utf-8")

        # ================================
        # Extract Text from PDF (PyMuPDF)
        # ================================
        elif filename.endswith(".pdf"):

            pdf_document = fitz.open(
                stream=file_bytes,
                filetype="pdf"
            )

            extracted_text = ""

            for page in pdf_document:

                page_text = page.get_text()

                extracted_text += page_text + "\n"

            pdf_document.close()

        # ================================
        # Extract Text from DOCX
        # ================================
        elif filename.endswith(".docx"):
            docx_file = DocxDocument(io.BytesIO(file_bytes))

            extracted_text = "\n".join(
                paragraph.text
                for paragraph in docx_file.paragraphs
            )

        else:
            raise HTTPException(
                status_code=400,
                detail="Only PDF, DOCX, and TXT files are supported."
            )

        # Prevent empty document upload
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from this document."
            )
        
        # ================================
        # Clean extracted text
        # Remove invalid Unicode characters
        # ================================
        extracted_text = "".join(
            char for char in extracted_text
            if not 0xD800 <= ord(char) <= 0xDFFF
        )

        extracted_text = (
            extracted_text
            .encode("utf-8", errors="ignore")
            .decode("utf-8", errors="ignore")
        )    
        # Generate embedding for document content
        document_embedding = generate_embedding(extracted_text)

        # Save document to PostgreSQL
        document = Document(
            user_id=current_user.id,
            filename=file.filename,
            content=extracted_text,
            embedding=document_embedding
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        print("DOCUMENT SAVED:", document.id, document.filename)

        return {
            "message": "Document uploaded successfully.",
            "filename": file.filename
        }

    except HTTPException:
        raise

    except Exception as e:
        print("UPLOAD ERROR:", str(e))
        return {
            "error": f"Upload failed: {str(e)}"
        }

# ================================
# 📄 Get Uploaded Documents
# ================================
@app.get("/documents")
def get_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    documents = db.query(Document).filter(
        Document.user_id == current_user.id
    ).order_by(
        Document.id.desc()
    ).all()

    return documents

# ================================
# 👁 Get Single Document
# ================================
@app.get("/documents/{document_id}")
def get_single_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()

    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found."
        )

    return document

# ================================
# 🗑 Delete Document
# ================================
@app.delete("/documents/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()

    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found."
        )

    db.delete(document)
    db.commit()

    return {
        "message": "Document deleted successfully."
    }    
# ================================
# Register User - PostgreSQL
# ================================
@app.post("/register")
def register(
    user: UserRegister,
    db: Session = Depends(get_db)
):

    # Check if username already exists
    existing_user = db.query(User).filter(
        User.username == user.username
    ).first()

    # Prevent duplicate usernames
    if existing_user:
        raise HTTPException(
            status_code=400,
            detail="Username already exists"
        )

    # Hash password before saving
    hashed_password = pwd_context.hash(user.password)

    # Create new user object
    new_user = User(
        username=user.username,
        hashed_password=hashed_password
    )

    # Save user into database
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    return {
        "message": "User registered successfully"
    }


# ================================
# Login User - PostgreSQL
# ================================
@app.post("/login")
def login(
    user: UserLogin,
    db: Session = Depends(get_db)
):

    # Find user in database
    db_user = db.query(User).filter(
        User.username == user.username
    ).first()

    # Check if user exists
    if not db_user:
        raise HTTPException(
            status_code=400,
            detail="Invalid username or password"
        )

    # Verify password
    valid_password = pwd_context.verify(
        user.password,
        db_user.hashed_password
    )

    # Wrong password
    if not valid_password:
        raise HTTPException(
            status_code=400,
            detail="Invalid username or password"
        )

    # Create JWT token
    token = jwt.encode(
        {"sub": db_user.username},
        SECRET_KEY,
        algorithm=ALGORITHM
    )

    return {
        "message": "Login successful",
        "token": token
    }

from datetime import datetime

